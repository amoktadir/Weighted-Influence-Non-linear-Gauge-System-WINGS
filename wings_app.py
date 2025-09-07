import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Circle
import graphviz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# Define linguistic terms mapping with abbreviations
LINGUISTIC_TERMS = {
    "strength": {
        "AS": 9,
        "ExS": 8,
        "VSS": 7,
        "SSP": 6,
        "SS": 5,
        "MSP": 4,
        "MS": 3,
        "WS": 2,
        "EqS": 1,
        "NS": 0
    },
    "influence": {
        "AI": 9,
        "ExI": 8,
        "VSI": 7,
        "SIP": 6,
        "SI": 5,
        "MIP": 4,
        "MI": 3,
        "WI": 2,
        "EqI": 1,
        "NI": 0
    }
}

# Full form mapping for tooltips
FULL_FORMS = {
    "AS": "Absolute strength",
    "ExS": "Extreme strength",
    "VSS": "Very strong strength",
    "SSP": "Strong strength plus",
    "SS": "Strong strength",
    "MSP": "Moderate strength plus",
    "MS": "Moderate strength",
    "WS": "Weak strength",
    "EqS": "Equal strength",
    "NS": "No strength",
    "AI": "Absolute influence",
    "ExI": "Extreme influence",
    "VSI": "Very strong influence",
    "SIP": "Strong influence plus",
    "SI": "Strong influence",
    "MIP": "Moderate influence plus",
    "MI": "Moderate influence",
    "WI": "Weak influence",
    "EqI": "Equal influence",
    "NI": "No influence"
}

def wings_method_experts(strengths_list, influence_matrices_list, weights=None):
    """
    Execute the WINGS method with multiple experts
    
    Parameters:
    strengths_list (list of lists): List of internal strength/importance from each expert
    influence_matrices_list (list of 2D arrays): List of influence matrices from each expert
    weights (list): List of weights for each expert (optional)
    
    Returns:
    dict: Dictionary containing all results including average matrices
    """
    n = len(strengths_list[0])  # Number of components
    num_experts = len(strengths_list)  # Number of experts
    
    # Create a list to store each expert's direct matrix D
    D_matrices = []
    
    # Construct the direct strength-influence matrix D for each expert
    for k in range(num_experts):
        D = np.array(influence_matrices_list[k], dtype=float)
        np.fill_diagonal(D, strengths_list[k])
        D_matrices.append(D)
    
    # Calculate the average direct matrix (weighted if weights are provided)
    if weights is not None:
        # Convert to numpy array for weighted average calculation
        D_arrays = np.array(D_matrices)
        D_avg = np.zeros_like(D_arrays[0])
        for i, weight in enumerate(weights):
            D_avg += weight * D_arrays[i]
        D_avg = D_avg / num_experts  # Divide by number of experts
    else:
        D_avg = np.mean(D_matrices, axis=0)
    
    # Step 2: Calculate the calibration factor s for the average matrix
    s = np.sum(D_avg)
    
    # Step 3: Calibrate the average matrix to get C
    C = D_avg / s
    
    # Step 4: Calculate the total strength-influence matrix T
    I = np.identity(n)
    try:
        T = np.linalg.inv(I - C) - I  # This is equivalent to C + C² + C³ + ...
    except np.linalg.LinAlgError:
        st.error("Matrix inversion failed. Please check your input values.")
        return None
    
    # Step 5: Calculate row sums (r) and column sums (c)
    r = np.sum(T, axis=1)
    c = np.sum(T, axis=0)
    
    # Step 6: Calculate prominence and relation
    prominence = r + c
    relation = r - c
    
    return {
        'expert_matrices': D_matrices,
        'average_direct_matrix': D_avg,
        'calibrated_matrix': C,
        'total_matrix': T,
        'row_sums': r,
        'column_sums': c,
        'prominence': prominence,
        'relation': relation
    }

def generate_flowchart_for_expert(expert_data, component_names, expert_idx=None, model_type="Linguistic Terms"):
    """
    Generate a flowchart for a specific expert based on their inputs
    
    Parameters:
    expert_data (dict): Dictionary containing expert inputs
    component_names (list): List of component names
    expert_idx (int): Index of the expert (for title)
    model_type (str): Type of model ("Linguistic Terms" or "Real Data")
    
    Returns:
    graphviz.Digraph: Flowchart graph for the expert
    """
    n = len(component_names)
    
    # Create the graph
    if expert_idx is not None:
        graph = graphviz.Digraph(comment=f'WINGS Analysis Flowchart - Expert {expert_idx+1}')
    else:
        graph = graphviz.Digraph(comment='WINGS Analysis Flowchart')
    
    graph.attr(rankdir='TD', size='8,8')
    
    # Add nodes with strength information
    for comp_idx, comp_name in enumerate(component_names):
        if model_type == "Linguistic Terms":
            strength = expert_data['strengths_linguistic'][comp_idx]
            label = f"{comp_name} ({strength})"
        else:
            strength_value = expert_data['strengths_numerical'][comp_idx]
            label = f"{comp_name} ({strength_value:.2f})"
        
        graph.node(comp_name, label=label, shape='box', style='rounded,filled', 
                  fillcolor='lightblue', fontsize='12')
    
    # Add edges with influence information
    for from_idx, from_comp in enumerate(component_names):
        for to_idx, to_comp in enumerate(component_names):
            if from_idx == to_idx:
                continue  # Skip self-influences
                
            if model_type == "Linguistic Terms":
                influence = expert_data['influence_matrix_linguistic'][from_idx][to_idx]
                if influence != "NI":  # Only add if there's an influence
                    graph.edge(from_comp, to_comp, label=influence)
            else:
                influence_value = expert_data['influence_matrix_numerical'][from_idx][to_idx]
                if influence_value != 0:  # Only add if there's an influence
                    graph.edge(from_comp, to_comp, label=f"{influence_value:.2f}")
    
    return graph

def create_word_report(results, component_names, model_type, n_experts=1, expert_weights=None):
    """
    Create a Word document report of the WINGS analysis
    
    Parameters:
    results (dict): Dictionary containing all results from WINGS analysis
    component_names (list): List of component names
    model_type (str): Type of model ("Linguistic Terms" or "Real Data")
    n_experts (int): Number of experts
    expert_weights (list): List of weights for each expert
    
    Returns:
    docx.Document: Word document with the report
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('WINGS Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date and model info
    from datetime import datetime
    date_para = doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    model_para = doc.add_paragraph(f"Model type: {model_type}")
    experts_para = doc.add_paragraph(f"Number of experts: {n_experts}")
    
    if expert_weights and n_experts > 1:
        weights_text = "Expert weights: " + ", ".join([f"Expert {i+1}: {weight:.2f}" for i, weight in enumerate(expert_weights)])
        doc.add_paragraph(weights_text)
    
    doc.add_paragraph()  # Add empty line
    
    # Add component names
    comp_para = doc.add_paragraph("Components analyzed: ")
    for i, name in enumerate(component_names):
        comp_para.add_run(f"{i+1}. {name}  ")
    
    doc.add_paragraph()  # Add empty line
    
    # Add prominence and relation results
    doc.add_heading('Prominence and Relation Results', level=1)
    
    # Create results table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Row Sum (r)'
    hdr_cells[2].text = 'Column Sum (c)'
    hdr_cells[3].text = 'Prominence (r+c)'
    hdr_cells[4].text = 'Relation (r-c)'
    
    # Add data rows
    for i, name in enumerate(component_names):
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = f"{results['row_sums'][i]:.3f}"
        row_cells[2].text = f"{results['column_sums'][i]:.3f}"
        row_cells[3].text = f"{results['prominence'][i]:.3f}"
        row_cells[4].text = f"{results['relation'][i]:.3f}"
    
    doc.add_paragraph()  # Add empty line
    
    # Add classification
    doc.add_heading('Component Classification', level=1)
    
    # Create classification table
    class_table = doc.add_table(rows=1, cols=3)
    class_table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = class_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Relation (r-c)'
    
    # Add data rows
    for i, name in enumerate(component_names):
        status = "Cause" if results['relation'][i] > 0 else "Effect"
        row_cells = class_table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = status
        row_cells[2].text = f"{results['relation'][i]:.3f}"
    
    doc.add_paragraph()  # Add empty line
    
    # Add matrices
    doc.add_heading('Matrices', level=1)
    
    # Average Direct Matrix
    doc.add_heading('Average Direct Matrix (D_avg)', level=2)
    df_D_avg = pd.DataFrame(results['average_direct_matrix'], index=component_names, columns=component_names)
    add_dataframe_to_doc(doc, df_D_avg)
    
    # Calibrated Matrix
    doc.add_heading('Calibrated Matrix (C)', level=2)
    df_C = pd.DataFrame(results['calibrated_matrix'], index=component_names, columns=component_names)
    add_dataframe_to_doc(doc, df_C)
    
    # Total Matrix
    doc.add_heading('Total Strength-Influence Matrix (T)', level=2)
    df_T = pd.DataFrame(results['total_matrix'], index=component_names, columns=component_names)
    add_dataframe_to_doc(doc, df_T)
    
    # Add interpretation
    doc.add_heading('Interpretation of Results', level=1)
    interpretation = doc.add_paragraph()
    interpretation.add_run("Prominence (r+c) ").bold = True
    interpretation.add_run("indicates the overall importance of a component in the system. Components with higher prominence values have greater overall impact.")
    
    interpretation = doc.add_paragraph()
    interpretation.add_run("Relation (r-c) ").bold = True
    interpretation.add_run("indicates whether a component is a cause or effect: ")
    interpretation.add_run("Positive values ").bold = True
    interpretation.add_run("indicate a component is a ")
    interpretation.add_run("Cause ").bold = True
    interpretation.add_run("(influences others more than it is influenced). ")
    interpretation.add_run("Negative values ").bold = True
    interpretation.add_run("indicate a component is an ")
    interpretation.add_run("Effect ").bold = True
    interpretation.add_run("(is influenced more than it influences others).")
    
    return doc

def add_dataframe_to_doc(doc, df, precision=3):
    """
    Add a pandas DataFrame to a Word document as a table
    
    Parameters:
    doc: Word document object
    df: pandas DataFrame to add
    precision: Number of decimal places to display
    """
    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns)+1)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''  # Empty cell for index
    
    for i, col in enumerate(df.columns):
        hdr_cells[i+1].text = str(col)
    
    # Add data rows
    for i, index in enumerate(df.index):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        
        for j, col in enumerate(df.columns):
            value = df.iloc[i, j]
            if isinstance(value, (int, float)):
                row_cells[j+1].text = f"{value:.{precision}f}"
            else:
                row_cells[j+1].text = str(value)
    
    doc.add_paragraph()  # Add empty line

def get_word_download_link(doc):
    """
    Generate a download link for a Word document
    
    Parameters:
    doc: Word document object
    
    Returns:
    str: HTML download link
    """
    # Save document to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Encode to base64
    b64 = base64.b64encode(file_stream.read()).decode()
    
    # Create download link
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="wings_analysis_report.docx">Download Word Report</a>'
    return href

def main():
    st.set_page_config(page_title="WINGS Method Analysis", layout="wide", page_icon="📊")
    st.title("📊 WINGS Method Analysis Platform")
    st.write("""
    This tool implements the Weighted Influence Non-linear Gauge System (WINGS) method 
    for analyzing systems with interrelated components, incorporating input from multiple experts/real data.
    """)
    
    # Add a new tab for How to Use
    tab_howto, tab_analysis = st.tabs(["📘 How to Use", "📊 Analysis"])
    
    with tab_howto:
        st.header("How to Use the WINGS Analysis Platform")
        
        st.markdown("""
        ### Overview
        The WINGS (Weighted Influence Non-linear Gauge System) method is a decision-making tool 
        that helps analyze complex systems with interrelated components. This platform allows 
        you to perform WINGS analysis using either linguistic terms or direct numerical values.
        
        ### Step-by-Step Guide
        
        1. **Configuration (Sidebar)**
           - Select your input model: Linguistic Terms or Real Data
           - Specify the number of components in your system
           - For linguistic models, specify the number of experts
           - Name each component for easy reference
        
        2. **Input Data**
           - **Component Strengths**: For each component, specify its internal strength/importance
           - **Influence Matrix**: Define how each component influences others
           - Use the expandable Linguistic Terms Mapping reference if needed
        
        3. **Run Analysis**
           - Click the "Run WINGS Analysis" button to process your inputs
           - The system will calculate prominence and relation values
        
        4. **Interpret Results**
           - **Flowchart**: Visual representation of components and their interactions
           - **Matrices**: View the various calculated matrices
           - **Results**: See prominence and relation values for each component
           - **Classification**: Components are classified as Causes or Effects
           - **Visualization**: Graphical representations of the analysis
        
        ### Input Models
        
        #### Linguistic Terms Model
        - Use predefined linguistic terms (e.g., AS, ExS, VSS for strength)
        - Multiple experts can provide assessments
        - Expert weights can be assigned for weighted averages
        - Terms are converted to numerical values for calculation
        
        #### Real Data Model
        - Input numerical values directly
        - Suitable when precise measurements are available
        - Single "expert" input mode
        
        ### Understanding the Results
        
        - **Prominence (r+c)**: Indicates the overall importance of a component in the system
        - **Relation (r-c)**: 
          - Positive values indicate a component is a **Cause** (influences others more than it's influenced)
          - Negative values indicate a component is an **Effect** (is influenced more than it influences others)
        
        ### Tips for Effective Use
        
        - Start with a small number of components to understand the method
        - Use descriptive names for components to make interpretation easier
        - For linguistic models, ensure all experts understand the term definitions
        - Review the flowchart to verify your inputs match your mental model of the system
        """)
        
        with st.expander("Linguistic Terms Reference"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**Strength Terms**")
                strength_data = []
                for abbr, value in LINGUISTIC_TERMS["strength"].items():
                    strength_data.append({
                        "Abbreviation": abbr,
                        "Full Form": FULL_FORMS[abbr],
                        "Numerical Value": value
                    })
                strength_df = pd.DataFrame(strength_data)
                st.dataframe(strength_df, hide_index=True, use_container_width=True)
            
            with col2:
                st.write("**Influence Terms**")
                influence_data = []
                for abbr, value in LINGUISTIC_TERMS["influence"].items():
                    influence_data.append({
                        "Abbreviation": abbr,
                        "Full Form": FULL_FORMS[abbr],
                        "Numerical Value": value
                    })
                influence_df = pd.DataFrame(influence_data)
                st.dataframe(influence_df, hide_index=True, use_container_width=True)
    
    with tab_analysis:
        # Model selection in sidebar
        with st.sidebar:
            st.header("⚙️ Configuration")
            model_type = st.selectbox(
                "Select Input Model",
                ["Linguistic Terms", "Real Data"],
                help="Choose between linguistic terms or direct numerical input"
            )
            
            n_components = st.number_input("Number of Components", min_value=2, max_value=25, value=3, help="How many components are in your system?")
            
            # Only show number of experts for linguistic model
            if model_type == "Linguistic Terms":
                n_experts = st.number_input("Number of Experts", min_value=1, max_value=15, value=1, help="How many experts will provide assessments?")
            else:
                n_experts = 1  # Real data model only has one "expert"
            
            component_names = []
            for i in range(n_components):
                name = st.text_input(f"Name of Component {i+1}", value=f"C{i+1}", key=f"comp_name_{i}")
                component_names.append(name)
            
            # Expert weights for linguistic model
            expert_weights = None
            if model_type == "Linguistic Terms" and n_experts > 1:
                st.markdown("---")
                st.subheader("Expert Weights")
                st.write("Assign weights to each expert (must sum to 1.0):")
                
                weights = []
                total_weight = 0
                
                for i in range(n_experts):
                    # Calculate the maximum allowed value for this weight
                    max_val = min(1.0, 1.0 - total_weight + (1.0/n_experts))
                    
                    weight = st.number_input(
                        f"Weight for Expert {i+1}", 
                        min_value=0.0, 
                        max_value=max_val,
                        value=1.0/n_experts,
                        step=0.01,
                        format="%.2f",
                        key=f"weight_{i}",
                        help=f"Maximum allowed: {max_val:.2f}"
                    )
                    weights.append(weight)
                    total_weight += weight
                
                # Display the current total
                st.write(f"**Current total:** {total_weight:.2f}/1.0")
                
                # Check if weights sum to 1
                if abs(total_weight - 1.0) > 0.001:
                    st.error(f"Weights must sum to 1.0. Current sum: {total_weight:.2f}")
                    st.stop()
                
                expert_weights = weights
                
            st.markdown("---")
            
            if model_type == "Linguistic Terms":
                st.info("💡 **Tip**: Use the abbreviations for strength and influence assessments.")
            else:
                st.info("💡 **Tip**: Enter numerical values directly (any real numbers).")
        
        # Display the linguistic terms mapping if selected
        if model_type == "Linguistic Terms":
            with st.expander("View Linguistic Terms Mapping", expanded=False):
                st.subheader("Linguistic Terms Mapping")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Strength Terms**")
                    strength_data = []
                    for abbr, value in LINGUISTIC_TERMS["strength"].items():
                        strength_data.append({
                            "Abbreviation": abbr,
                            "Full Form": FULL_FORMS[abbr],
                            "Numerical Value": value
                        })
                    strength_df = pd.DataFrame(strength_data)
                    st.dataframe(strength_df, hide_index=True, use_container_width=True)
                
                with col2:
                    st.write("**Influence Terms**")
                    influence_data = []
                    for abbr, value in LINGUISTIC_TERMS["influence"].items():
                        influence_data.append({
                            "Abbreviation": abbr,
                            "Full Form": FULL_FORMS[abbr],
                            "Numerical Value": value
                        })
                    influence_df = pd.DataFrame(influence_data)
                    st.dataframe(influence_df, hide_index=True, use_container_width=True)
        
        # Initialize session state for experts data
        if 'experts_data' not in st.session_state:
            st.session_state.experts_data = {}
        
        # Initialize or update session state for each expert
        for expert_idx in range(n_experts):
            if expert_idx not in st.session_state.experts_data:
                # Initialize with default values
                if model_type == "Linguistic Terms":
                    st.session_state.experts_data[expert_idx] = {
                        'strengths_linguistic': ["SS" if i == 0 else "MS" for i in range(n_components)],
                        'strengths_numerical': [5.0 if i == 0 else 3.0 for i in range(n_components)],
                        'influence_matrix_linguistic': [[list(LINGUISTIC_TERMS["influence"].keys())[0] for _ in range(n_components)] for _ in range(n_components)],
                        'influence_matrix_numerical': [[0.0 for _ in range(n_components)] for _ in range(n_components)]
                    }
                else:
                    st.session_state.experts_data[expert_idx] = {
                        'strengths_numerical': [5.0 if i == 0 else 3.0 for i in range(n_components)],
                        'influence_matrix_numerical': [[0.0 for _ in range(n_components)] for _ in range(n_components)]
                    }
            else:
                # Ensure the data structure matches the current number of components
                if len(st.session_state.experts_data[expert_idx]['strengths_numerical']) != n_components:
                    if model_type == "Linguistic Terms":
                        st.session_state.experts_data[expert_idx]['strengths_linguistic'] = ["SS" if i == 0 else "MS" for i in range(n_components)]
                    st.session_state.experts_data[expert_idx]['strengths_numerical'] = [5.0 if i == 0 else 3.0 for i in range(n_components)]
                
                if (len(st.session_state.experts_data[expert_idx]['influence_matrix_numerical']) != n_components or 
                    any(len(row) != n_components for row in st.session_state.experts_data[expert_idx]['influence_matrix_numerical'])):
                    if model_type == "Linguistic Terms":
                        st.session_state.experts_data[expert_idx]['influence_matrix_linguistic'] = [[list(LINGUISTIC_TERMS["influence"].keys())[0] for _ in range(n_components)] for _ in range(n_components)]
                    st.session_state.experts_data[expert_idx]['influence_matrix_numerical'] = [[0.0 for _ in range(n_components)] for _ in range(n_components)]
        
        # Main content area
        if model_type == "Linguistic Terms" and n_experts > 1:
            st.header(f"👨‍💼 Expert Input ({n_experts} Experts)")
            
            # Create tabs for each expert
            expert_tabs = st.tabs([f"Expert {i+1}" for i in range(n_experts)])
        else:
            st.header("📊 Data Input")
            expert_tabs = [st.container()]  # Single container for real data model
        
        strengths_list = []
        influence_matrices_list = []
        
        for expert_idx in range(n_experts):
            if model_type == "Linguistic Terms" and n_experts > 1:
                tab = expert_tabs[expert_idx]
            else:
                tab = expert_tabs[0]
                
            with tab:
                if model_type == "Linguistic Terms" and n_experts > 1:
                    st.subheader(f"Expert {expert_idx+1} Input")
                    if expert_weights:
                        st.write(f"**Weight:** {expert_weights[expert_idx]:.2f}")
                elif model_type == "Real Data":
                    st.subheader("Direct Numerical Input")
                
                # Component strengths for this expert
                st.write("**Component Strengths**")
                strengths = []
                
                # Create a table for strength inputs
                strength_cols = st.columns(n_components + 1)
                with strength_cols[0]:
                    st.markdown("**Component**")
                for i in range(n_components):
                    with strength_cols[i + 1]:
                        st.markdown(f"**{component_names[i]}**")
                
                strength_input_cols = st.columns(n_components + 1)
                with strength_input_cols[0]:
                    st.markdown("**Strength Value**")
                
                for i in range(n_components):
                    with strength_input_cols[i + 1]:
                        if model_type == "Linguistic Terms":
                            # Get the current value from session state
                            current_strength = st.session_state.experts_data[expert_idx]['strengths_linguistic'][i]
                            
                            strength_term = st.selectbox(
                                f"Strength of {component_names[i]}", 
                                options=list(LINGUISTIC_TERMS["strength"].keys()),
                                index=list(LINGUISTIC_TERMS["strength"].keys()).index(current_strength),
                                key=f"strength_{expert_idx}_{i}",
                                help=FULL_FORMS[current_strength],
                                label_visibility="collapsed"
                            )
                            
                            # Convert to numerical value
                            strength_value = LINGUISTIC_TERMS["strength"][strength_term]
                            
                            # Display the numerical value
                            st.markdown(f"**{strength_value}**")
                            
                            # Update session state
                            st.session_state.experts_data[expert_idx]['strengths_linguistic'][i] = strength_term
                        else:
                            # Direct numerical input
                            strength_value = st.number_input(
                                f"Strength of {component_names[i]}", 
                                value=float(st.session_state.experts_data[expert_idx]['strengths_numerical'][i]),
                                key=f"strength_{expert_idx}_{i}",
                                label_visibility="collapsed",
                                step=0.1,
                                format="%.2f"
                            )
                        
                        strengths.append(strength_value)
                        st.session_state.experts_data[expert_idx]['strengths_numerical'][i] = strength_value
                
                # Influence matrix for this expert
                st.write("**Influence Matrix**")
                st.write("Enter the influence between components (row influences column):")
                
                influence_matrix = np.zeros((n_components, n_components))
                
                # Create a grid for the influence matrix
                for i in range(n_components):
                    st.markdown(f"**Influences from {component_names[i]}**")
                    
                    # Create header row
                    header_cols = st.columns(n_components + 1)
                    with header_cols[0]:
                        st.markdown("**To →**")
                    for j in range(n_components):
                        with header_cols[j + 1]:
                            st.markdown(f"**{component_names[j]}**")
                    
                    # Create input row
                    input_cols = st.columns(n_components + 1)
                    with input_cols[0]:
                        st.markdown(f"**From {component_names[i]}**")
                    for j in range(n_components):
                        with input_cols[j + 1]:
                            if i == j:
                                st.markdown("—", help="Diagonal elements represent self-strength")
                            else:
                                if model_type == "Linguistic Terms":
                                    # Get the current value from session state
                                    current_influence = st.session_state.experts_data[expert_idx]['influence_matrix_linguistic'][i][j]
                                    
                                    influence_term = st.selectbox(
                                        f"{component_names[i]} → {component_names[j]}", 
                                        options=list(LINGUISTIC_TERMS["influence"].keys()),
                                        index=list(LINGUISTIC_TERMS["influence"].keys()).index(current_influence),
                                        key=f"inf_{expert_idx}_{i}_{j}",
                                        label_visibility="collapsed",
                                        help=FULL_FORMS[current_influence]
                                    )
                                    
                                    # Convert to numerical value
                                    influence_value = LINGUISTIC_TERMS["influence"][influence_term]
                                    
                                    # Display the numerical value
                                    st.markdown(f"**{influence_value}**")
                                    
                                    # Update session state
                                    st.session_state.experts_data[expert_idx]['influence_matrix_linguistic'][i][j] = influence_term
                                else:
                                    # Direct numerical input
                                    influence_value = st.number_input(
                                        f"{component_names[i]} → {component_names[j]}", 
                                        value=float(st.session_state.experts_data[expert_idx]['influence_matrix_numerical'][i][j]),
                                        key=f"inf_{expert_idx}_{i}_{j}",
                                        label_visibility="collapsed",
                                        step=0.1,
                                        format="%.2f"
                                    )
                                    st.markdown(f"**{influence_value}**")
                                
                                influence_matrix[i][j] = influence_value
                                st.session_state.experts_data[expert_idx]['influence_matrix_numerical'][i][j] = influence_value
            
            strengths_list.append(strengths)
            influence_matrices_list.append(influence_matrix.tolist())
        
        # Execute analysis when button is clicked
        if st.button("🚀 Run WINGS Analysis", type="primary", use_container_width=True):
            with st.spinner("Calculating..."):
                results = wings_method_experts(strengths_list, influence_matrices_list, expert_weights)
            
            if results is None:
                return
            
            st.success("Analysis Complete!")
            
            # Display results in tabs - Flowchart first
            tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
                "🔗 Flowchart", "📋 Expert Matrices", "🧮 Average Matrix", 
                "📊 Results", "🏷️ Component Classification", "📈 Visualization", "📤 Export"
            ])
            
            with tab1:
                st.subheader("Component Interaction Flowchart")
                
                # Generate flowcharts based on the model type
                if model_type == "Linguistic Terms" and n_experts > 1:
                    # Generate a flowchart for each expert
                    for expert_idx in range(n_experts):
                        st.subheader(f"Flowchart for Expert {expert_idx+1}")
                        flowchart = generate_flowchart_for_expert(
                            st.session_state.experts_data[expert_idx], 
                            component_names,
                            expert_idx,
                            model_type
                        )
                        st.graphviz_chart(flowchart, use_container_width=True)
                else:
                    # Generate a single flowchart (for single expert or real data)
                    if model_type == "Linguistic Terms":
                        expert_idx = 0
                    else:
                        expert_idx = None  # Real data
                    
                    flowchart = generate_flowchart_for_expert(
                        st.session_state.experts_data[0], 
                        component_names,
                        expert_idx,
                        model_type
                    )
                    st.graphviz_chart(flowchart, use_container_width=True)
                
                # Add explanation
                if model_type == "Linguistic Terms":
                    st.markdown("""
                    **Flowchart Explanation:**
                    - **Nodes**: Represent components with their strength level in parentheses
                    - **Edges**: Show influences between components with their influence level
                    - **Strength Levels**: AS (Absolute), ExS (Extreme), VSS (Very Strong), SSP (Strong Plus), 
                      SS (Strong), MSP (Moderate Plus), MS (Moderate), WS (Weak), EqS (Equal), NS (No Strength)
                    - **Influence Levels**: AI (Absolute), ExI (Extreme), VSI (Very Strong), SIP (Strong Plus), 
                      SI (Strong), MIP (Moderate Plus), MI (Moderate), WI (Weak), EqI (Equal), NI (No Influence)
                    """)
                else:
                    st.markdown("""
                    **Flowchart Explanation:**
                    - **Nodes**: Represent components with their strength value in parentheses
                    - **Edges**: Show influences between components with their influence value
                    - **Values**: Direct numerical inputs as provided
                    """)
            
            with tab2:
                if model_type == "Linguistic Terms" and n_experts > 1:
                    st.subheader("Individual Expert Matrices")
                    for i, D in enumerate(results['expert_matrices']):
                        st.write(f"**Expert {i+1} Direct Matrix D:**")
                        df_D = pd.DataFrame(D, index=component_names, columns=component_names)
                        st.dataframe(df_D.style.format("{:.3f}"), use_container_width=True)
                else:
                    st.subheader("Direct Matrix D")
                    df_D = pd.DataFrame(results['expert_matrices'][0], index=component_names, columns=component_names)
                    st.dataframe(df_D.style.format("{:.3f}"), use_container_width=True)
            
            with tab3:
                st.subheader("Average Direct Matrix D_avg")
                df_D_avg = pd.DataFrame(
                    results['average_direct_matrix'], 
                    index=component_names, 
                    columns=component_names
                )
                st.dataframe(df_D_avg.style.format("{:.3f}"), use_container_width=True)
                
                st.subheader("Calibrated Matrix C")
                df_C = pd.DataFrame(
                    results['calibrated_matrix'], 
                    index=component_names, 
                    columns=component_names
                )
                st.dataframe(df_C.style.format("{:.3f}"), use_container_width=True)
                
                st.subheader("Total Strength-Influence Matrix T")
                df_T = pd.DataFrame(
                    results['total_matrix'], 
                    index=component_names, 
                    columns=component_names
                )
                st.dataframe(df_T.style.format("{:.3f}"), use_container_width=True)
            
            with tab4:
                st.subheader("Prominence and Relation Values")
                results_df = pd.DataFrame({
                    'Component': component_names,
                    'Row Sum (r)': results['row_sums'],
                    'Column Sum (c)': results['column_sums'],
                    'Prominence (r+c)': results['prominence'],
                    'Relation (r-c)': results['relation']
                }).sort_values('Prominence (r+c)', ascending=False)
                
                # Format the DataFrame
                styled_df = results_df.style.format({
                    'Row Sum (r)': '{:.3f}',
                    'Column Sum (c)': '{:.3f}',
                    'Prominence (r+c)': '{:.3f}',
                    'Relation (r-c)': '{:.3f}'
                }).background_gradient(subset=['Prominence (r+c)'], cmap='Blues')
                
                st.dataframe(styled_df, use_container_width=True, hide_index=True)
            
            with tab5:
                st.subheader("Component Classification")
                
                # Create a DataFrame for better visualization
                classification_data = []
                for i, (name, rel, prom) in enumerate(zip(component_names, results['relation'], results['prominence'])):
                    status = "Cause" if rel > 0 else "Effect"
                    classification_data.append({
                        "Component": name,
                        "Type": status,
                        "Relation (r-c)": rel,
                        "Prominence (r+c)": prom
                    })
                
                classification_df = pd.DataFrame(classification_data)
                
                # Display as metrics
                cols = st.columns(3)
                for i, row in classification_df.iterrows():
                    with cols[i % 3]:
                        emoji = "➡️" if row['Type'] == 'Cause' else "⬅️"
                        st.metric(
                            label=f"{emoji} {row['Component']}",
                            value=row['Type'],
                            delta=f"Relation: {row['Relation (r-c)']:.3f}"
                        )
                
                # Display as bar chart
                fig, ax = plt.subplots(figsize=(10, 6))
                colors = ['#2ecc71' if t == 'Cause' else '#e74c3c' for t in classification_df['Type']]
                bars = ax.bar(classification_df['Component'], classification_df['Relation (r-c)'], color=colors)
                ax.set_title('Component Relation Values')
                ax.set_ylabel('Relation (r-c)')
                plt.xticks(rotation=45)
                st.pyplot(fig)
            
            with tab6:
                st.subheader("Visualization")
                
                # Create a scatter plot of prominence vs relation
                fig, ax = plt.subplots(figsize=(10, 8))
                
                # Plot components
                for i, name in enumerate(component_names):
                    color = 'green' if results['relation'][i] > 0 else 'red'
                    ax.scatter(results['prominence'][i], results['relation'][i], s=150, color=color, alpha=0.7)
                    ax.annotate(name, (results['prominence'][i], results['relation'][i]), 
                               xytext=(5, 5), textcoords='offset points', fontsize=12)
                
                # Add reference lines
                ax.axhline(y=0, color='gray', linestyle='--', alpha=0.7)
                ax.axvline(x=np.mean(results['prominence']), color='gray', linestyle='--', alpha=0.7)
                
                # Labels and title
                ax.set_xlabel('Prominence (r+c)')
                ax.set_ylabel('Relation (r-c)')
                ax.set_title('Component Analysis: Prominence vs Relation')
                ax.grid(True, alpha=0.3)
                
                st.pyplot(fig)
                
                # Bar chart of prominence
                fig2, ax2 = plt.subplots(figsize=(10, 6))
                y_pos = np.arange(len(component_names))
                ax2.barh(y_pos, results['prominence'], alpha=0.7)
                ax2.set_yticks(y_pos)
                ax2.set_yticklabels(component_names)
                ax2.set_xlabel('Prominence (r+c)')
                ax2.set_title('Component Prominence Ranking')
                ax2.grid(True, alpha=0.3, axis='x')
                
                st.pyplot(fig2)
                
                # Create a cause-effect quad chart
                st.subheader("Cause-Effect Diagram")
                avg_prominence = np.mean(results['prominence'])
                
                # Create matplotlib figure for more control
                fig3, ax3 = plt.subplots(figsize=(10, 8))
                
                # Define quadrants
                ax3.axhline(y=0, color='gray', linestyle='--', alpha=0.7)
                ax3.axvline(x=avg_prominence, color='gray', linestyle='--', alpha=0.7)
                
                # Plot points
                for i, name in enumerate(component_names):
                    x = results['prominence'][i]
                    y = results['relation'][i]
                    color = 'green' if y > 0 else 'red'
                    ax3.scatter(x, y, s=150, color=color, alpha=0.7)
                    ax3.annotate(name, (x, y), xytext=(5, 5), textcoords='offset points', fontsize=12)
                
                # Add quadrant labels
                ax3.text(0.02, 0.98, "Cause Components", transform=ax3.transAxes, fontsize=14, 
                       verticalalignment='top', bbox=dict(boxstyle='round', facecolor='green', alpha=0.2))
                ax3.text(0.02, 0.02, "Effect Components", transform=ax3.transAxes, fontsize=14, 
                       verticalalignment='bottom', bbox=dict(boxstyle='round', facecolor='red', alpha=0.2))
                
                # Set labels and title
                ax3.set_xlabel('Prominence (r+c)')
                ax3.set_ylabel('Relation (r-c)')
                ax3.set_title('Cause-Effect Diagram')
                ax3.grid(True, alpha=0.3)
                
                st.pyplot(fig3)
            
            with tab7:
                st.subheader("Export Results")
                st.write("Download a comprehensive report of your WINGS analysis in Word format.")
                
                # Create Word document
                doc = create_word_report(results, component_names, model_type, n_experts, expert_weights)
                
                # Generate download link
                html_link = get_word_download_link(doc)
                st.markdown(html_link, unsafe_allow_html=True)
                
                st.info("The Word report includes:")
                st.markdown("""
                - Analysis parameters and configuration
                - Prominence and relation results table
                - Component classification table
                - All calculated matrices (D_avg, C, T)
                - Interpretation of results
                """)

if __name__ == "__main__":
    main()
